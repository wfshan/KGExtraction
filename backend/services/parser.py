"""
文档解析器
支持 PDF、TXT、Markdown、DOCX（含表格）、CSV、XLSX 格式的文本提取。
结构化文件（CSV/XLSX）统一输出为「表头行 + 每行 col | col | col」，
首个非空行即表头，供输入理解阶段作为候选 Schema 信号。
"""
import csv
from pathlib import Path


def parse_document(file_path: Path, file_type: str) -> str:
    """
    解析文档，提取纯文本内容

    Args:
        file_path: 文件路径
        file_type: 文件类型 (pdf/txt/md/docx/csv/xlsx)

    Returns:
        提取的纯文本
    """
    parsers = {
        "pdf": _parse_pdf,
        "txt": _parse_txt,
        "md": _parse_markdown,
        "docx": _parse_docx,
        "csv": _parse_csv,
        "xlsx": _parse_xlsx,
    }

    parser = parsers.get(file_type)
    if not parser:
        raise ValueError(f"不支持的文件类型: {file_type}")

    text = parser(file_path)

    # 基础清洗
    text = _clean_text(text)
    return text


def _parse_pdf(file_path: Path) -> str:
    """解析 PDF 文件"""
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    texts = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            texts.append(page_text)
    return "\n\n".join(texts)


def _parse_txt(file_path: Path) -> str:
    """解析 TXT 文件"""
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def _parse_markdown(file_path: Path) -> str:
    """解析 Markdown 文件，剥离行内标记但保留标题的 "#" 前缀。

    标题标记是层次切分（hierarchical）识别层级的依据；此前被一并剥离，
    导致解析后的纯文本无法再按标题层级切分。
    """
    from markdown_it import MarkdownIt

    with open(file_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    md = MarkdownIt()
    tokens = md.parse(md_text)

    texts = []
    heading_level = 0
    for token in tokens:
        if token.type == "heading_open":
            tag = token.tag or "h1"
            heading_level = int(tag[1]) if len(tag) > 1 and tag[1].isdigit() else 1
            continue
        if token.type == "heading_close":
            heading_level = 0
            continue
        if token.children:
            content = "".join(
                child.content for child in token.children
                if child.type in ("text", "code_inline")
            )
            if not content:
                continue
            if heading_level and content.strip():
                texts.append(f"{'#' * heading_level} {content.strip()}")
            else:
                texts.append(content)
        elif token.content:
            texts.append(token.content)

    return "\n".join(texts)


def _docx_heading_level(paragraph) -> int:
    """识别 DOCX 段落的标题层级（Heading 1-6 / 标题 1-6），非标题返回 0。"""
    import re
    try:
        style_name = paragraph.style.name or ""
    except Exception:
        return 0
    m = re.match(r"^(?:Heading|标题)\s*([1-6])$", style_name.strip(), re.IGNORECASE)
    return int(m.group(1)) if m else 0


def _parse_docx(file_path: Path) -> str:
    """解析 DOCX 文件。

    标题样式（Heading/标题 1-6）转为 Markdown "#" 前缀保留在纯文本中，
    使层次切分（hierarchical）对 DOCX 也能按真实标题层级工作，而非退化为普通切分。
    """
    from docx import Document

    doc = Document(str(file_path))
    texts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            level = _docx_heading_level(paragraph)
            if level:
                texts.append(f"{'#' * level} {paragraph.text.strip()}")
            else:
                texts.append(paragraph.text)

    # 表格内容
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                texts.append(" | ".join(row_texts))

    return "\n\n".join(texts)


def _parse_csv(file_path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            with open(file_path, "r", encoding=encoding, newline="") as f:
                reader = csv.reader(f)
                rows = []
                for row in reader:
                    cells = [cell.strip() for cell in row]
                    if any(cells):
                        rows.append(" | ".join(cells))
                return "\n".join(rows)
        except UnicodeDecodeError:
            continue
    raise ValueError("CSV 文件编码不受支持，请使用 UTF-8 或 GB18030 编码")


def _parse_xlsx(file_path: Path) -> str:
    """解析 Excel（.xlsx）。多工作表分别输出，每表首行为表头。"""
    from openpyxl import load_workbook

    wb = load_workbook(str(file_path), read_only=True, data_only=True)
    sheets = []
    for ws in wb.worksheets:
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [("" if v is None else str(v)).strip() for v in row]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            # 多表时标注表名，便于后续识别表间关系
            header = f"# 工作表: {ws.title}" if len(wb.worksheets) > 1 else ""
            sheets.append((header + "\n" if header else "") + "\n".join(rows))
    wb.close()
    return "\n\n".join(sheets)


def _clean_text(text: str) -> str:
    """基础文本清洗"""
    import re

    # 去除多余空行
    text = re.sub(r"\n{3,}", "\n\n", text)
    # 去除行首尾空白
    lines = [line.strip() for line in text.split("\n")]
    text = "\n".join(lines)
    # 去除首尾空白
    text = text.strip()

    return text
